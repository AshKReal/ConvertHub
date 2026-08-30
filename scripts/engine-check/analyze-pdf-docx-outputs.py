"""specs/000-engine-quality-check.md - "exit code 0" isn't quality. This inspects what each PDF->DOCX
candidate's output docx actually contains: extractable paragraph text, real Word tables (not frozen
frames/images), and inline shapes (a proxy for "content became a picture instead of text"). Compares
the three candidates' outputs for the same input files.

Usage: python analyze-pdf-docx-outputs.py
"""
import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from lib import HERE

CANDIDATES = {
    "A (libreoffice)": HERE / "results" / "pdf-docx-a-out",
    "B (pdf2docx)": HERE / "results" / "pdf-docx-b-out",
    "C (aspose)": HERE / "results" / "pdf-docx-c-out",
}


def inspect(path: Path) -> dict:
    doc = Document(str(path))
    text_len = sum(len(p.text) for p in doc.paragraphs)
    table_count = len(doc.tables)
    table_cell_text_len = sum(len(c.text) for t in doc.tables for row in t.rows for c in row.cells)
    # inline/floating pictures: a proxy for "PDF content became an image, not text"
    picture_count = len(doc.element.body.findall(f".//{qn('pic:pic')}"))
    return {
        "text_len": text_len,
        "table_count": table_count,
        "table_cell_text_len": table_cell_text_len,
        "picture_count": picture_count,
    }


results = {}
for label, out_dir in CANDIDATES.items():
    files = sorted(out_dir.glob("*.docx")) if out_dir.exists() else []
    rows = []
    for f in files:
        try:
            rows.append({"file": f.name, **inspect(f)})
        except Exception as exc:
            rows.append({"file": f.name, "error": str(exc)})
    results[label] = rows

out_path = HERE / "results" / "pdf-docx-comparison.json"
out_path.write_text(json.dumps(results, indent=2), encoding="utf-8")

print(f"{'file':<35}" + "".join(f"{label:>20}" for label in CANDIDATES))
all_files = sorted({r["file"] for rows in results.values() for r in rows})
for file in all_files:
    line = f"{file:<35}"
    for label in CANDIDATES:
        row = next((r for r in results[label] if r["file"] == file), None)
        if row is None or "error" in row:
            line += f"{'MISSING/ERR':>20}"
        else:
            line += f"{row['text_len']}c/{row['table_count']}t/{row['picture_count']}p".rjust(20)
    print(line)

print("\ncolumns: extractable-text-chars / real-tables / inline-pictures")
print(f"Full comparison: {out_path}")
